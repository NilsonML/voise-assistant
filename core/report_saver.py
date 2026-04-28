# core/report_saver.py
# Сохраняем диалоги в файлы: TXT и/или DOCX

import os
from datetime import datetime
from docx import Document
from config import Config
from core.logger import log


class ReportSaver:
    """Умеет сохранять распознанный текст и ответ ИИ в файл"""
    
    def __init__(self):
        self.default_format = "both"   # txt, docx, both, none
        self.output_dir = Config.OUTPUT_DIR   # Папка для сохранения
    
    def save_report(self, text, analysis=None, audio_file=None, video_file=None, output_dir=None, format=None):
        """
        Сохраняет отчёт.
        
        text — распознанный текст пользователя
        analysis — ответ ИИ
        audio_file — откуда взято аудио (для информации в отчёте)
        video_file — откуда взято видео
        output_dir — куда сохранять (если не указано, берём из Config)
        format — 'txt', 'docx', 'both', 'none'
        """
        output_dir = output_dir or self.output_dir
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        save_format = format or self.default_format
        
        # Готовим содержимое отчёта
        content = self._prepare_content(text, analysis, audio_file, video_file)
        full_text = '\n'.join(content)
        
        results = {'txt': None, 'docx': None}
        
        # Сохраняем TXT
        if save_format in ['txt', 'both']:
            txt_path = os.path.join(output_dir, f"report_{timestamp}.txt")
            try:
                with open(txt_path, 'w', encoding='utf-8') as f:
                    f.write(full_text)
                log.info(f"TXT отчет сохранен: {txt_path}")
                results['txt'] = txt_path
            except Exception as e:
                log.error(f"Ошибка сохранения TXT: {e}")
        
        # Сохраняем DOCX (красивый форматированный документ)
        if save_format in ['docx', 'both']:
            docx_path = os.path.join(output_dir, f"report_{timestamp}.docx")
            try:
                doc = self._create_docx(content)
                doc.save(docx_path)
                log.info(f"DOCX отчет сохранен: {docx_path}")
                results['docx'] = docx_path
            except Exception as e:
                log.error(f"Ошибка сохранения DOCX: {e}")
        
        return results
    
    def _prepare_content(self, text, analysis, audio_file, video_file):
        """Формируем список строк для отчёта (общий для TXT и DOCX)"""
        content = []
        content.append("=" * 60)
        content.append("ГОЛОСОВОЙ АССИСТЕНТ - ОТЧЕТ")
        content.append("=" * 60)
        content.append(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M:%S')}")
        content.append("=" * 60)
        content.append("")
        
        # Откуда взялся текст
        if video_file:
            content.append(f"📹 Источник: Видео файл - {os.path.basename(video_file)}")
            content.append("")
        elif audio_file:
            content.append(f"🎵 Источник: Аудио файл - {os.path.basename(audio_file)}")
            content.append("")
            
        content.append("📝 РАСПОЗНАННЫЙ ТЕКСТ:")
        content.append("-" * 40)
        content.append(text)
        content.append("")
        
        if analysis:
            content.append("🤖 АНАЛИЗ / ОТВЕТ ИИ:")
            content.append("-" * 40)
            content.append(analysis)
            content.append("")
            
        content.append("=" * 60)
        
        return content
    
    def _create_docx(self, content_lines):
        """Создаёт DOCX документ с заголовками и параграфами"""
        doc = Document()
        doc.add_heading('Голосовой Ассистент - Отчет', 0)
        
        for line in content_lines:
            # Пропускаем разделительные линии (они в DOCX не нужны)
            if line.startswith('=') or line.startswith('-'):
                continue
            # Заголовки
            elif line.startswith('📝') or line.startswith('🤖'):
                doc.add_heading(line, level=1)
            # Обычный текст
            elif line.strip():
                doc.add_paragraph(line)
        
        return doc